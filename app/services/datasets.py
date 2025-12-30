from __future__ import annotations

import csv
import hashlib
import io
import json
import random
import uuid
from dataclasses import dataclass
from datetime import datetime, timezone
from pathlib import Path
from typing import Any, Iterable, Optional

import numpy as np
from sklearn.model_selection import train_test_split

from app.core.config import Settings
from app.db.metadata import insert_dataset
from app.schemas.dataset import DatasetUploadOptions
from app.services.jsonl import read_jsonl, write_jsonl


CANONICAL_REQUIRED_FIELDS = {"id", "text", "label", "source"}


@dataclass(frozen=True)
class DatasetPaths:
    root_dir: Path
    meta_path: Path
    train_path: Path
    eval_path: Path
    holdout_path: Path


def _utc_now_iso() -> str:
    return datetime.now(timezone.utc).replace(microsecond=0).isoformat()


def _sha256_bytes(data: bytes) -> str:
    h = hashlib.sha256()
    h.update(data)
    return h.hexdigest()


def _normalize_label(raw: Any, opts: DatasetUploadOptions) -> Optional[int]:
    if raw is None:
        return None

    if isinstance(raw, bool):
        return int(raw)
    if isinstance(raw, int) and raw in {0, 1}:
        return int(raw)
    if isinstance(raw, float) and raw in {0.0, 1.0}:
        return int(raw)
    s = str(raw).strip().lower()
    if s in {v.lower() for v in opts.legit_label_values}:
        return 0
    if s in {v.lower() for v in opts.scam_label_values}:
        return 1

    if s in {"0", "1"}:
        return int(s)

    try:
        f = float(s)
        if f in {0.0, 1.0}:
            return int(f)
    except Exception:
        pass

    return None


def _detect_text_column(fieldnames: list[str], opts: DatasetUploadOptions) -> Optional[str]:
    lowered = {c.lower(): c for c in fieldnames}

    text_candidates = ["text", "message", "sms", "content", "body"]

    if opts.text_column:
        return opts.text_column

    for k in text_candidates:
        if k in lowered:
            return lowered[k]

    return None


def _detect_label_column(fieldnames: list[str], opts: DatasetUploadOptions) -> Optional[str]:
    lowered = {c.lower(): c for c in fieldnames}
    label_candidates = [
        "label",
        "class",
        "category",
        "target",
        "is_scam",
        "is_spam",
        "is_fraudulent",
        "is_fraud",
        "fraud",
    ]

    if opts.label_column:
        return opts.label_column

    for k in label_candidates:
        if k in lowered:
            return lowered[k]

    return None


def _detect_columns(fieldnames: list[str], opts: DatasetUploadOptions) -> tuple[str, str]:
    text_col = _detect_text_column(fieldnames, opts)
    label_col = _detect_label_column(fieldnames, opts)

    if not text_col or not label_col:
        raise ValueError(
            f"Could not detect columns. Found columns={fieldnames}. "
            f"Provide text_column and label_column explicitly."
        )

    return text_col, label_col


def _parse_sms_spam_collection(raw_text: str) -> list[dict[str, Any]]:
    rows: list[dict[str, Any]] = []
    for line in raw_text.splitlines():
        line = line.strip("\n")
        if not line:
            continue
        if "\t" not in line:
            continue
        label, text = line.split("\t", 1)
        label = label.strip().lower()
        if label not in {"ham", "spam", "smish", "smishing"}:
            continue
        rows.append({"label": 0 if label == "ham" else 1, "text": text.strip()})
    return rows


def _build_structured_text(values: dict[str, Any], exclude_keys: set[str]) -> str:
    parts: list[str] = []
    for k in sorted(values.keys()):
        if k in exclude_keys:
            continue
        v = values.get(k)
        if v is None:
            continue
        s = str(v).strip()
        if not s:
            continue
        parts.append(f"{k}={s}")
    return " | ".join(parts)


def _parse_csv(raw_text: str, opts: DatasetUploadOptions) -> list[dict[str, Any]]:
    reader = csv.DictReader(io.StringIO(raw_text))
    if not reader.fieldnames:
        return []

    text_col = _detect_text_column(reader.fieldnames, opts)
    label_col = _detect_label_column(reader.fieldnames, opts)
    if label_col is None and opts.default_label is None:
        raise ValueError(
            f"Could not detect label column. Found columns={reader.fieldnames}. "
            f"Provide label_column explicitly or set default_label."
        )

    out: list[dict[str, Any]] = []
    for row in reader:
        if text_col is not None:
            text = (row.get(text_col) or "").strip()
        else:
            exclude = {k for k in [label_col] if k is not None}
            text = _build_structured_text(row, exclude)
        if not text:
            continue

        if label_col is None:
            label = opts.default_label
        else:
            label = _normalize_label(row.get(label_col), opts)
        if label is None:
            label = opts.default_label
        if label not in {0, 1}:
            continue

        exclude_meta = {k for k in [text_col, label_col] if k is not None}
        metadata = {k: v for k, v in row.items() if k not in exclude_meta}
        out.append({"label": int(label), "text": text, "metadata": metadata})

    return out


def _parse_xlsx(file_bytes: bytes, opts: DatasetUploadOptions) -> list[dict[str, Any]]:
    try:
        from openpyxl import load_workbook
    except Exception as e:
        raise ValueError(f"Missing dependency for .xlsx support: {e}")

    wb = load_workbook(io.BytesIO(file_bytes), read_only=True, data_only=True)
    ws = wb.active

    rows_iter = ws.iter_rows(values_only=True)
    try:
        header = next(rows_iter)
    except StopIteration:
        return []

    fieldnames: list[str] = []
    for idx, v in enumerate(header):
        name = str(v).strip() if v is not None else ""
        if not name:
            name = f"col_{idx}"
        fieldnames.append(name)

    text_col = _detect_text_column(fieldnames, opts)
    label_col = _detect_label_column(fieldnames, opts)
    if label_col is None and opts.default_label is None:
        raise ValueError(
            f"Could not detect label column. Found columns={fieldnames}. "
            f"Provide label_column explicitly or set default_label."
        )

    out: list[dict[str, Any]] = []
    for row in rows_iter:
        values = {fieldnames[i]: row[i] if i < len(row) else None for i in range(len(fieldnames))}
        if text_col is not None:
            text = str(values.get(text_col) or "").strip()
        else:
            exclude = {k for k in [label_col] if k is not None}
            text = _build_structured_text(values, exclude)
        if not text:
            continue

        if label_col is None:
            label = opts.default_label
        else:
            label = _normalize_label(values.get(label_col), opts)
        if label is None:
            label = opts.default_label
        if label not in {0, 1}:
            continue

        exclude_meta = {k for k in [text_col, label_col] if k is not None}
        metadata = {k: v for k, v in values.items() if k not in exclude_meta}
        out.append({"label": int(label), "text": text, "metadata": metadata})

    return out


def _parse_docx(file_bytes: bytes, opts: DatasetUploadOptions) -> list[dict[str, Any]]:
    if opts.default_label not in {0, 1}:
        raise ValueError(".docx upload requires default_label=0 or default_label=1")

    try:
        from docx import Document
    except Exception as e:
        raise ValueError(f"Missing dependency for .docx support: {e}")

    doc = Document(io.BytesIO(file_bytes))

    texts: list[str] = []
    for p in doc.paragraphs:
        t = (p.text or "").strip()
        if t:
            texts.append(t)

    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if (c.text or "").strip()]
            if cells:
                texts.append(" | ".join(cells))

    out: list[dict[str, Any]] = []
    seen: set[str] = set()
    for t in texts:
        if t in seen:
            continue
        seen.add(t)
        out.append({"label": int(opts.default_label), "text": t, "metadata": {"source": "docx"}})

    return out


def _parse_json(raw_text: str, opts: DatasetUploadOptions) -> list[dict[str, Any]]:
    rows: list[dict[str, Any]] = []

    lines = [ln for ln in raw_text.splitlines() if ln.strip()]
    if lines and all(ln.lstrip().startswith("{") for ln in lines[:5]):
        for ln in lines:
            try:
                obj = json.loads(ln)
            except json.JSONDecodeError:
                rows = []
                break
            rows.append(obj)
        if rows:
            return _normalize_json_rows(rows, opts)

    try:
        obj = json.loads(raw_text)
    except json.JSONDecodeError:
        return []

    if isinstance(obj, list):
        return _normalize_json_rows(obj, opts)

    if isinstance(obj, dict) and "data" in obj and isinstance(obj["data"], list):
        return _normalize_json_rows(obj["data"], opts)

    return []


def _normalize_json_rows(rows: Iterable[dict[str, Any]], opts: DatasetUploadOptions) -> list[dict[str, Any]]:
    out: list[dict[str, Any]] = []
    for row in rows:
        if not isinstance(row, dict):
            continue

        if "text" in row and "label" in row:
            text = str(row.get("text") or "").strip()
            label = _normalize_label(row.get("label"), opts)
            if text and label is not None:
                out.append({"text": text, "label": int(label), "metadata": row.get("metadata") or {}})
            continue

        text = None
        for k in ("TEXT", "text", "message", "sms", "content"):
            if k in row and row[k] is not None:
                text = str(row[k]).strip()
                break

        label_raw = None
        for k in ("LABEL", "label", "class", "category", "is_scam", "is_spam", "is_fraudulent", "is_fraud"):
            if k in row:
                label_raw = row.get(k)
                break

        if not text:
            exclude_keys = {
                "TEXT",
                "text",
                "message",
                "sms",
                "content",
                "LABEL",
                "label",
                "class",
                "category",
                "is_scam",
                "is_spam",
                "is_fraudulent",
                "is_fraud",
            }
            text = _build_structured_text(row, exclude_keys)
        if not text:
            continue
        label = _normalize_label(label_raw, opts)
        if label is None:
            label = opts.default_label
        if label not in {0, 1}:
            continue

        exclude_meta = {
            "TEXT",
            "text",
            "message",
            "sms",
            "content",
            "LABEL",
            "label",
            "class",
            "category",
            "is_scam",
            "is_spam",
            "is_fraudulent",
            "is_fraud",
        }
        metadata = {k: v for k, v in row.items() if k not in exclude_meta}
        out.append({"text": text, "label": int(label), "metadata": metadata})

    return out


def _balance(rows: list[dict[str, Any]], strategy: str, seed: int) -> list[dict[str, Any]]:
    if strategy == "none":
        return rows
    if strategy != "downsample_majority":
        return rows

    scams = [r for r in rows if r.get("label") == 1]
    legit = [r for r in rows if r.get("label") == 0]

    if not scams or not legit:
        return rows

    rnd = random.Random(seed)

    if len(scams) > len(legit):
        scams = rnd.sample(scams, k=len(legit))
    else:
        legit = rnd.sample(legit, k=len(scams))

    balanced = scams + legit
    rnd.shuffle(balanced)
    return balanced


def _split(rows: list[dict[str, Any]], opts: DatasetUploadOptions) -> tuple[list[dict[str, Any]], list[dict[str, Any]], list[dict[str, Any]]]:
    if not np.isclose(opts.train_ratio + opts.eval_ratio + opts.holdout_ratio, 1.0):
        raise ValueError("train_ratio + eval_ratio + holdout_ratio must sum to 1.0")

    if opts.eval_ratio == 0 and opts.holdout_ratio == 0:
        return rows, [], []

    if opts.train_ratio == 1.0:
        return rows, [], []

    y = [int(r["label"]) for r in rows]
    counts = {0: y.count(0), 1: y.count(1)}
    can_stratify_first = len(set(y)) > 1 and min(counts.values()) >= 2

    train_rows, tmp_rows = train_test_split(
        rows,
        test_size=(1.0 - opts.train_ratio),
        random_state=opts.seed,
        stratify=y if can_stratify_first else None,
    )

    tmp_ratio = opts.eval_ratio + opts.holdout_ratio
    if tmp_ratio <= 0:
        return train_rows, [], []

    if opts.holdout_ratio == 0:
        return train_rows, tmp_rows, []

    if opts.eval_ratio == 0:
        return train_rows, [], tmp_rows

    eval_share = opts.eval_ratio / tmp_ratio

    y_tmp = [int(r["label"]) for r in tmp_rows]
    counts_tmp = {0: y_tmp.count(0), 1: y_tmp.count(1)}
    can_stratify_second = len(set(y_tmp)) > 1 and min(counts_tmp.values()) >= 2
    eval_rows, holdout_rows = train_test_split(
        tmp_rows,
        test_size=(1.0 - eval_share),
        random_state=opts.seed,
        stratify=y_tmp if can_stratify_second else None,
    )

    return train_rows, eval_rows, holdout_rows


def _canonicalize_rows(rows: list[dict[str, Any]], source: str, channel: str) -> list[dict[str, Any]]:
    out: list[dict[str, Any]] = []
    for r in rows:
        text = str(r.get("text") or "").strip()
        label = r.get("label")
        if not text or label not in {0, 1}:
            continue
        sample = {
            "id": str(uuid.uuid4()),
            "text": text,
            "label": int(label),
            "source": source,
            "channel": channel,
            "metadata": r.get("metadata") or {},
        }
        out.append(sample)
    return out


def dataset_paths(settings: Settings, dataset_id: str) -> DatasetPaths:
    root = settings.datasets_dir / dataset_id
    return DatasetPaths(
        root_dir=root,
        meta_path=root / "meta.json",
        train_path=root / "train.jsonl",
        eval_path=root / "eval.jsonl",
        holdout_path=root / "holdout.jsonl",
    )


def ingest_dataset(
    *,
    settings: Settings,
    original_filename: str,
    file_bytes: bytes,
    opts: DatasetUploadOptions,
) -> dict[str, Any]:
    dataset_id = str(uuid.uuid4())
    version = "v1"
    created_at = _utc_now_iso()

    raw_hash = _sha256_bytes(file_bytes)
    raw_dir = settings.raw_datasets_dir / dataset_id
    raw_dir.mkdir(parents=True, exist_ok=True)
    raw_path = raw_dir / original_filename
    raw_path.write_bytes(file_bytes)

    parsed: list[dict[str, Any]] = []
    lower_name = original_filename.lower()

    if lower_name.endswith(".xlsx"):
        parsed = _parse_xlsx(file_bytes, opts)
        source = "public_dataset"
        label_mapping = {"legit_values": opts.legit_label_values, "scam_values": opts.scam_label_values}
    elif lower_name.endswith(".docx"):
        parsed = _parse_docx(file_bytes, opts)
        source = "public_dataset"
        label_mapping = {"default_label": int(opts.default_label) if opts.default_label is not None else None}
    else:
        text = file_bytes.decode("utf-8", errors="replace")

        if lower_name == "smsspamcollection" or "smsspamcollection" in lower_name:
            parsed = _parse_sms_spam_collection(text)
            source = "public_dataset"
            label_mapping = {"ham": 0, "spam": 1}
        elif lower_name.endswith(".txt"):
            parsed = _parse_sms_spam_collection(text)
            if not parsed:
                raise ValueError("Unsupported .txt format. Expected tab-separated label and text")
            source = "public_dataset"
            label_mapping = {"ham": 0, "spam": 1, "smish": 1, "smishing": 1}
        elif lower_name.endswith(".csv"):
            parsed = _parse_csv(text, opts)
            source = "public_dataset"
            label_mapping = {"legit_values": opts.legit_label_values, "scam_values": opts.scam_label_values}
        elif lower_name.endswith(".json") or lower_name.endswith(".jsonl"):
            parsed = _parse_json(text, opts)
            source = "public_dataset"
            label_mapping = {"legit_values": opts.legit_label_values, "scam_values": opts.scam_label_values}
        else:
            raise ValueError(
                "Unsupported file type. Upload .csv, .json/.jsonl, .txt, .xlsx, .docx, or SMSSpamCollection"
            )

    parsed = _balance(parsed, opts.balance_strategy, opts.seed)
    canonical = _canonicalize_rows(parsed, source, opts.channel)

    if not canonical:
        raise ValueError("No valid samples parsed from dataset")

    train_rows, eval_rows, holdout_rows = _split(canonical, opts)

    paths = dataset_paths(settings, dataset_id)
    paths.root_dir.mkdir(parents=True, exist_ok=True)
    write_jsonl(paths.train_path, train_rows)
    write_jsonl(paths.eval_path, eval_rows)
    write_jsonl(paths.holdout_path, holdout_rows)

    num_scam = sum(1 for r in canonical if r["label"] == 1)
    num_legit = sum(1 for r in canonical if r["label"] == 0)

    meta = {
        "dataset_id": dataset_id,
        "version": version,
        "created_at": created_at,
        "dataset_name": opts.dataset_name,
        "source_filename": original_filename,
        "raw_sha256": raw_hash,
        "num_samples": len(canonical),
        "num_scam": num_scam,
        "num_legit": num_legit,
        "splits": {"train": len(train_rows), "eval": len(eval_rows), "holdout": len(holdout_rows)},
        "label_mapping": label_mapping,
        "split_config": {
            "train_ratio": opts.train_ratio,
            "eval_ratio": opts.eval_ratio,
            "holdout_ratio": opts.holdout_ratio,
            "seed": opts.seed,
            "balance_strategy": opts.balance_strategy,
        },
    }

    paths.meta_path.write_text(json.dumps(meta, ensure_ascii=False, indent=2), encoding="utf-8")

    insert_dataset(
        settings=settings,
        dataset_id=dataset_id,
        version=version,
        created_at=created_at,
        source_filename=original_filename,
        num_samples=len(canonical),
        num_scam=num_scam,
        num_legit=num_legit,
    )

    return meta


def load_dataset_meta(settings: Settings, dataset_id: str) -> dict[str, Any]:
    paths = dataset_paths(settings, dataset_id)
    if not paths.meta_path.exists():
        raise FileNotFoundError(f"Unknown dataset_id={dataset_id}")
    return json.loads(paths.meta_path.read_text(encoding="utf-8"))


def create_augmented_dataset(
    *,
    settings: Settings,
    base_dataset_id: str,
    added_train_rows: list[dict[str, Any]],
    dataset_name: str | None = None,
    seed: int = 1337,
) -> dict[str, Any]:
    base_meta = load_dataset_meta(settings, base_dataset_id)
    base_paths = dataset_paths(settings, base_dataset_id)

    base_train = list(read_jsonl(base_paths.train_path))
    base_eval = list(read_jsonl(base_paths.eval_path))
    base_holdout = list(read_jsonl(base_paths.holdout_path))

    dataset_id = str(uuid.uuid4())
    version = "v1"
    created_at = _utc_now_iso()

    paths = dataset_paths(settings, dataset_id)
    paths.root_dir.mkdir(parents=True, exist_ok=True)

    train_rows = base_train + list(added_train_rows)
    random.Random(seed).shuffle(train_rows)

    write_jsonl(paths.train_path, train_rows)
    write_jsonl(paths.eval_path, base_eval)
    write_jsonl(paths.holdout_path, base_holdout)

    all_rows = train_rows + base_eval + base_holdout
    num_scam = sum(1 for r in all_rows if int(r.get("label", 0)) == 1)
    num_legit = sum(1 for r in all_rows if int(r.get("label", 0)) == 0)

    meta = {
        "dataset_id": dataset_id,
        "version": version,
        "created_at": created_at,
        "dataset_name": dataset_name or base_meta.get("dataset_name") or f"augmented_{base_dataset_id}",
        "source_filename": base_meta.get("source_filename") or "augmented_dataset",
        "raw_sha256": base_meta.get("raw_sha256"),
        "parent_dataset_id": base_dataset_id,
        "augmentation": {
            "num_added_train": int(len(added_train_rows)),
        },
        "num_samples": int(len(all_rows)),
        "num_scam": int(num_scam),
        "num_legit": int(num_legit),
        "splits": {"train": len(train_rows), "eval": len(base_eval), "holdout": len(base_holdout)},
        "label_mapping": base_meta.get("label_mapping") or {},
        "split_config": base_meta.get("split_config") or {},
    }

    paths.meta_path.write_text(json.dumps(meta, ensure_ascii=False, indent=2), encoding="utf-8")

    insert_dataset(
        settings=settings,
        dataset_id=dataset_id,
        version=version,
        created_at=created_at,
        source_filename=str(meta.get("source_filename") or "augmented_dataset"),
        num_samples=int(len(all_rows)),
        num_scam=int(num_scam),
        num_legit=int(num_legit),
    )

    return meta
